from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\LENOVO\AppData\Local\Temp\informe_devops_wuepa_formulas_dora_source.docx"
)
OUTPUT = Path(
    r"C:\Users\LENOVO\OneDrive\Escritorio\wuepa\frontend_wuepa\informe_devops_wuepa_formulas_dora_apa.docx"
)

FORMULAS = [
    {
        "prefix": "Frecuencia de despliegue (DF):",
        "name": "Frecuencia de despliegue (DF)",
        "formula": "DF = N / T",
        "number": 1,
        "definition": (
            "Donde N es el número de despliegues exitosos en producción y T es la duración "
            "del periodo evaluado."
        ),
    },
    {
        "prefix": "Tiempo de espera para cambios (LT):",
        "name": "Tiempo de espera para cambios (LT)",
        "formula": "LT = (1 / M) × Σ(i=1 a M) [t_producción,i - t_commit,i]",
        "number": 2,
        "definition": (
            "Donde M es el número de cambios desplegados; t_producción,i es la fecha y hora "
            "en que el cambio i llegó correctamente a producción, y t_commit,i es la fecha y "
            "hora del primer commit asociado a ese cambio."
        ),
    },
    {
        "prefix": "Tasa de fallos en cambios (CFR):",
        "name": "Tasa de fallos en cambios (CFR)",
        "formula": "CFR = (D_fallidos / D_totales) × 100 %",
        "number": 3,
        "definition": (
            "Donde D_fallidos es la cantidad de despliegues que provocaron incidentes, "
            "reversiones o intervenciones, y D_totales es la cantidad total de despliegues "
            "realizados en producción durante el mismo periodo."
        ),
    },
    {
        "prefix": "Tiempo medio de restauración (MTTR):",
        "name": "Tiempo medio de restauración (MTTR)",
        "formula": "MTTR = (1 / K) × Σ(j=1 a K) [t_resolución,j - t_detección,j]",
        "number": 4,
        "definition": (
            "Donde K es el número de incidentes; t_resolución,j es la fecha y hora en que el "
            "incidente j quedó resuelto por completo, y t_detección,j es la fecha y hora en "
            "que dicho incidente fue detectado."
        ),
    },
]


def set_font(run, italic=False, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.italic = italic
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def insert_after(paragraph: Paragraph, style=None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if style is not None:
        created.style = style
    return created


def format_body_paragraph(paragraph: Paragraph):
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def make_equation_paragraph(after: Paragraph, formula: str, number: int) -> Paragraph:
    paragraph = insert_after(after, "Normal")
    format_body_paragraph(paragraph)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.25), WD_TAB_ALIGNMENT.CENTER
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
    )

    paragraph.add_run("\t")
    formula_run = paragraph.add_run(formula)
    set_font(formula_run, italic=True)
    paragraph.add_run("\t")
    number_run = paragraph.add_run(f"({number})")
    set_font(number_run)
    return paragraph


def format_formula_block(paragraph: Paragraph, data: dict):
    paragraph.clear()
    paragraph.style = "Body Text"
    format_body_paragraph(paragraph)
    paragraph.paragraph_format.keep_with_next = True

    name_run = paragraph.add_run(data["name"])
    set_font(name_run, bold=True)
    rest_run = paragraph.add_run(f" se calcula mediante la Ecuación {data['number']}.")
    set_font(rest_run)

    equation = make_equation_paragraph(
        paragraph, data["formula"], data["number"]
    )
    equation.paragraph_format.keep_with_next = True

    definition = insert_after(equation, "Body Text")
    format_body_paragraph(definition)
    definition.paragraph_format.keep_together = True
    definition_run = definition.add_run(data["definition"])
    set_font(definition_run)


def format_guide_citation(paragraph: Paragraph):
    paragraph.clear()
    paragraph.style = "Body Text"
    format_body_paragraph(paragraph)
    before = paragraph.add_run(
        "Las métricas DORA permiten evaluar el desempeño de entrega mediante frecuencia de "
        "despliegue, tiempo de espera para cambios, tasa de fallos en cambios y tiempo medio de "
        "restauración. Para que los resultados futuros sean comparables, Wuepa deberá calcularlas "
        "con las fórmulas de la guía proporcionada ("
    )
    set_font(before)
    title = paragraph.add_run("Guía de implementación")
    set_font(title, italic=True)
    after = paragraph.add_run(
        ", s. f.). Como el pipeline actual no registra despliegues ni incidentes de producción, "
        "todavía no existen datos suficientes para obtener valores observados."
    )
    set_font(after)


def format_guide_reference(paragraph: Paragraph):
    paragraph.clear()
    paragraph.style = "Body Text"
    format_body_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    title = paragraph.add_run("Guía de implementación: Métricas DevOps (DORA).")
    set_font(title, italic=True)
    date = paragraph.add_run(" (s. f.).")
    set_font(date)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    document = Document(SOURCE)

    context = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Las métricas DORA permiten evaluar")
    )
    format_guide_citation(context)

    for data in FORMULAS:
        paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.startswith(data["prefix"])
        )
        format_formula_block(paragraph, data)

    reference = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Guía de implementación: Métricas DevOps (DORA).")
    )
    format_guide_reference(reference)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
