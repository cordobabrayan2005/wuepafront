from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(
    r"C:\Users\LENOVO\AppData\Local\Temp\informe_devops_wuepa_formulas_dora_apa_source.docx"
)
OUTPUT = Path(
    r"C:\Users\LENOVO\OneDrive\Escritorio\wuepa\frontend_wuepa\informe_devops_wuepa_dora_apa_aplicabilidad.docx"
)


def set_font(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")


def format_apa_paragraph(paragraph: Paragraph):
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def insert_after(paragraph: Paragraph, style=None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if style is not None:
        created.style = style
    return created


def add_metric_reason(after: Paragraph, metric: str, reason: str) -> Paragraph:
    paragraph = insert_after(after, "Body Text")
    format_apa_paragraph(paragraph)
    paragraph.paragraph_format.keep_together = True
    name = paragraph.add_run(f"{metric}. ")
    set_font(name, bold=True)
    text = paragraph.add_run(reason)
    set_font(text)
    return paragraph


def replace_cell(cell, text: str):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(text)
    set_font(run)


def main():
    document = Document(SOURCE)

    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "Estas fórmulas no deben aplicarse a las ejecuciones de compilación"
        )
    )

    heading = insert_after(anchor, "Heading 1")
    heading.add_run("Aplicabilidad actual de las métricas DORA en Wuepa")

    intro = insert_after(heading, "Body Text")
    format_apa_paragraph(intro)
    intro_run = intro.add_run(
        "Aunque las fórmulas anteriores son metodológicamente correctas, actualmente no pueden "
        "aplicarse para calcular resultados reales de Wuepa. El workflow de GitHub Actions se limita "
        "a instalar dependencias, compilar el frontend y guardar la carpeta dist como artefacto. No "
        "incluye una etapa de despliegue a producción ni un sistema de monitoreo de incidentes."
    )
    set_font(intro_run)

    cursor = add_metric_reason(
        intro,
        "Frecuencia de despliegue (DF)",
        "No puede calcularse porque el pipeline no publica versiones en producción y, por tanto, "
        "no existe un historial verificable de despliegues exitosos ni un periodo de medición asociado.",
    )
    cursor = add_metric_reason(
        cursor,
        "Tiempo de espera para cambios (LT)",
        "No puede calcularse porque no se registra la fecha y hora en que cada commit llega a "
        "producción. La hora de compilación del workflow no equivale a la hora de despliegue.",
    )
    cursor = add_metric_reason(
        cursor,
        "Tasa de fallos en cambios (CFR)",
        "No puede calcularse porque no hay despliegues de producción registrados ni una relación "
        "entre esos despliegues y posibles incidentes, reversiones o correcciones urgentes.",
    )
    cursor = add_metric_reason(
        cursor,
        "Tiempo medio de restauración (MTTR)",
        "No puede calcularse porque Wuepa no cuenta, dentro del alcance revisado, con alertas y "
        "registros que indiquen cuándo se detectó un incidente y cuándo se restauró el servicio.",
    )

    conclusion = insert_after(cursor, "Body Text")
    format_apa_paragraph(conclusion)
    conclusion.paragraph_format.keep_with_next = True
    conclusion_run = conclusion.add_run(
        "En consecuencia, las métricas DORA se incluyen como marco de medición futura y no como "
        "resultados observados. Una ejecución exitosa de compilación demuestra integración continua, "
        "pero no debe contabilizarse como un despliegue en producción."
    )
    set_font(conclusion_run)

    table_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "Línea base estimada y objetivos DORA para Wuepa"
    )
    table_heading.text = "Estado actual y objetivos DORA para Wuepa"

    table = document.tables[1]
    replace_cell(table.rows[0].cells[1], "Estado actual")
    statuses = [
        "Sin medición: no hay despliegues registrados",
        "Sin medición: no hay marcas de producción",
        "Sin medición: no hay despliegues e incidentes vinculados",
        "Sin medición: no hay alertas ni tiempos de resolución",
    ]
    for row, status in zip(table.rows[1:], statuses):
        replace_cell(row.cells[1], status)

    note = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nota. Los valores son estimaciones académicas")
    )
    note.clear()
    note.style = "Normal"
    format_apa_paragraph(note)
    label = note.add_run("Nota. ")
    set_font(label)
    label.font.italic = True
    body = note.add_run(
        "Wuepa no dispone actualmente de una línea base DORA cuantitativa. Los objetivos son "
        "referencias académicas y solo podrán compararse con resultados reales cuando existan "
        "despliegues trazables y registros de incidentes de producción."
    )
    set_font(body)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
