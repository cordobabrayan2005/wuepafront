from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\LENOVO\Downloads\informe_devops_wuepa_con_evidencias.docx")
OUTPUT = Path(
    r"C:\Users\LENOVO\OneDrive\Escritorio\wuepa\frontend_wuepa\informe_devops_wuepa_formulas_dora.docx"
)


def insert_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if style is not None:
        created.style = style
    if text:
        created.add_run(text)
    return created


def add_formula(after: Paragraph, name: str, formula: str, explanation: str) -> Paragraph:
    paragraph = insert_after(after, style="Body Text")
    paragraph.paragraph_format.keep_together = True
    name_run = paragraph.add_run(f"{name}: ")
    name_run.bold = True
    formula_run = paragraph.add_run(formula)
    formula_run.bold = True
    formula_run.italic = True
    paragraph.add_run(f". {explanation}")
    return paragraph


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = text
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    document = Document(SOURCE)

    dora_context = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Las métricas DORA permiten evaluar")
    )
    dora_context.text = (
        "Las métricas DORA permiten evaluar el desempeño de entrega mediante frecuencia de "
        "despliegue, tiempo de espera para cambios, tasa de fallos en cambios y tiempo medio de "
        "restauración. Para que los resultados futuros sean comparables, Wuepa deberá calcularlas "
        "con las fórmulas de la guía proporcionada (Guía de implementación, s. f.). Como el pipeline "
        "actual no registra despliegues ni incidentes de producción, todavía no existen datos "
        "suficientes para obtener valores observados."
    )
    dora_context.style = document.styles["Body Text"]

    heading = insert_after(dora_context, "Fórmulas DORA aplicables a Wuepa", "Heading 1")
    intro = insert_after(
        heading,
        "Las variables deberán obtenerse para un mismo periodo de evaluación y conservar una unidad "
        "de tiempo uniforme, por ejemplo, horas o días.",
        "Body Text",
    )

    cursor = add_formula(
        intro,
        "Frecuencia de despliegue (DF)",
        "DF = N / T",
        "N es el número de despliegues exitosos en producción y T es la duración del periodo evaluado.",
    )
    cursor = add_formula(
        cursor,
        "Tiempo de espera para cambios (LT)",
        "LT = (1 / M) × Σ(i=1 a M) [t_producción,i - t_commit,i]",
        "M es el número de cambios desplegados; para cada cambio se resta la fecha y hora del commit "
        "a la fecha y hora en que llegó correctamente a producción, y después se calcula el promedio.",
    )
    cursor = add_formula(
        cursor,
        "Tasa de fallos en cambios (CFR)",
        "CFR = (D_fallidos / D_totales) × 100 %",
        "D_fallidos representa los despliegues que provocaron incidentes, rollback o intervención, y "
        "D_totales representa todos los despliegues de producción del mismo periodo.",
    )
    cursor = add_formula(
        cursor,
        "Tiempo medio de restauración (MTTR)",
        "MTTR = (1 / K) × Σ(j=1 a K) [t_resolución,j - t_detección,j]",
        "K es el número de incidentes; para cada incidente se mide el tiempo desde su detección hasta "
        "la restauración completa del servicio y luego se obtiene el promedio.",
    )
    closing = insert_after(
        cursor,
        "Estas fórmulas no deben aplicarse a las ejecuciones de compilación del pipeline como si fueran "
        "despliegues: solo cuentan eventos e incidentes reales del entorno de producción.",
        "Body Text",
    )
    closing.paragraph_format.keep_with_next = True

    dora_table = document.tables[1]
    updates = [
        ("Frecuencia de despliegue (DF)", "N despliegues exitosos y periodo T"),
        ("Tiempo de espera para cambios (LT)", "M cambios y marcas de commit/producción"),
        ("Tasa de fallos en cambios (CFR)", "D_fallidos y D_totales del periodo"),
        ("Tiempo medio de restauración (MTTR)", "K incidentes y marcas de detección/resolución"),
    ]
    for row, (metric, source) in zip(dora_table.rows[1:], updates):
        replace_cell_text(row.cells[0], metric)
        replace_cell_text(row.cells[3], source)

    table_note = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nota. Los valores son estimaciones académicas")
    )
    table_note.text = (
        "Nota. Los valores son estimaciones académicas, no estadísticas observadas del proyecto. "
        "Deberán reemplazarse por resultados calculados con las fórmulas anteriores y con datos de "
        "GitHub, del proveedor de despliegue y del sistema de monitoreo."
    )
    table_note.style = document.styles["Normal"]

    vite_reference = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Vite. (2026). Building for production")
    )
    guide_reference = insert_after(vite_reference, style="Body Text")
    title_run = guide_reference.add_run("Guía de implementación: Métricas DevOps (DORA).")
    title_run.italic = True
    guide_reference.add_run(" (s. f.).")
    guide_reference.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
