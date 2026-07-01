from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


SOURCE = Path(r"C:\Users\LENOVO\Downloads\informe_devops_wuepa_normas_apa.docx")
OUTPUT = Path(r"C:\Users\LENOVO\OneDrive\Escritorio\wuepa\frontend_wuepa\informe_devops_wuepa_con_evidencias.docx")

FIGURES = [
    (
        Path(r"C:\Users\LENOVO\AppData\Local\Temp\codex-clipboard-edd32e70-e155-47ca-8375-853fbb5da54b.png"),
        "Resumen de la ejecución exitosa del pipeline",
        "Captura de GitHub Actions que muestra el estado exitoso, la duración y el artefacto frontend-wuepa-dist generado por el pipeline.",
        "Resumen de ejecución exitosa del pipeline de GitHub Actions",
    ),
    (
        Path(r"C:\Users\LENOVO\AppData\Local\Temp\codex-clipboard-b05991b6-c3bb-4d4a-8bad-7bd449684c21.png"),
        "Detalle de los pasos ejecutados por el pipeline",
        "La ejecución completó correctamente la descarga del código, la configuración de Node.js, la instalación de dependencias, la compilación y el almacenamiento del artefacto.",
        "Detalle de pasos completados por el pipeline de GitHub Actions",
    ),
    (
        Path(r"C:\Users\LENOVO\AppData\Local\Temp\codex-clipboard-5042444f-6dea-4109-944e-6d5e9fa58222.png"),
        "Archivo de configuración del workflow ejecutado",
        "Captura del archivo .github/workflows/pipeline.yml utilizado en la ejecución del pipeline del frontend Wuepa.",
        "Archivo pipeline.yml mostrado en GitHub Actions",
    ),
]


def insert_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if style is not None:
        created.style = style
    if text:
        created.add_run(text)
    return created


def set_alt_text(run, description: str) -> None:
    drawings = run._element.xpath(".//wp:docPr")
    if drawings:
        drawings[0].set("descr", description)


def add_figure(after: Paragraph, number: int, figure) -> Paragraph:
    image_path, title, note, alt_text = figure

    label = insert_after(after, f"Figura {number}", "Heading 1")

    title_paragraph = insert_after(label, style="Normal")
    title_run = title_paragraph.add_run(title)
    title_run.italic = True
    title_paragraph.paragraph_format.keep_with_next = True

    image_paragraph = insert_after(title_paragraph, style="Normal")
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(image_path), width=Inches(6.35))
    set_alt_text(image_run, alt_text)

    note_paragraph = insert_after(image_paragraph, style="Normal")
    note_label = note_paragraph.add_run("Nota. ")
    note_label.italic = True
    note_paragraph.add_run(note)
    return note_paragraph


def main() -> None:
    for image_path, *_ in FIGURES:
        if not image_path.exists():
            raise FileNotFoundError(image_path)

    document = Document(SOURCE)
    placeholder = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Para la entrega final se debe insertar aquí")
    )

    placeholder.text = ""
    placeholder.style = document.styles["Normal"]
    cursor = placeholder

    for index, figure in enumerate(FIGURES, start=2):
        if index > 2:
            page_break = insert_after(cursor, style="Normal")
            page_break.add_run().add_break(WD_BREAK.PAGE)
            cursor = page_break
        cursor = add_figure(cursor, index, figure)

    final_page_break = insert_after(cursor, style="Normal")
    final_page_break.add_run().add_break(WD_BREAK.PAGE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
